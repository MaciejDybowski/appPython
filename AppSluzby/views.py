from django.db.models.functions import Coalesce
from django.http import HttpResponse
from django.shortcuts import render, redirect

# Create your views here.
from django.shortcuts import render
from django.views.generic import TemplateView
from .models import PersonOnDuty,PersonOnCeremony,Person,Duty,Ceremony,Soldier
from .forms import PersonForm, DutyForm, PersonOnDutyForm, CeremonyForm, PersonOnCeremonyForm


from docx import Document
from docx.shared import Pt,Cm,Inches
import json
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date,timedelta



# Create your views here.

def counter(type):  # funkcja typu switch case bo nie pythonie nie ma domyslnie switcha - wykorzystywana do obliczania
                    # sluzb dla kazej osoby
    switcher = {
        'Kmp': 1,
        'Sto': 0.5,
        'Str': 1,
        'Pst': 1,
        'Bat': 1,
        'War': 1
    }
    return switcher.get(type, "Invalid input")


class HomePageView(TemplateView):
    def get(self, request, **kwargs):
        import xlrd
        """loc = ("Pluton.xlsx")
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        rows = sheet.nrows
        i = 0
        while i < rows:
            degreeTemp = sheet.cell_value(i,0)
            nameTemp = sheet.cell_value(i, 1)
            surnameTemp = sheet.cell_value(i, 2)
            dutyTemp = sheet.cell_value(i,3)
            ceremonyTemp = sheet.cell_value(i,4)
            Person.objects.create(degree=degreeTemp, name=nameTemp,surname=surnameTemp,numberOfDuty=dutyTemp,
                                  numberOfCeremony=ceremonyTemp)
            i = i + 1


        loc = ("Dane.xlsx")
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)

        i = 0
        while i < sheet.nrows:
            tSoldier = Soldier()
            tSoldier.degree = sheet.cell_value(i, 0)
            tSoldier.name = sheet.cell_value(i, 1)
            tSoldier.surname = sheet.cell_value(i, 2)
            tSoldier.save()
            i = i + 1"""

        # pobieranie listy osob i liczenie ile maja sluzb
        personList = Person.objects.all()
        for i in personList:
            numberOfDuty = 0.0
            Dutys = PersonOnDuty.objects.filter(idPerson=i.idPerson)
            for p in Dutys:
                numberOfDuty += counter(p.idDuty.typeOfDuty)
            i.numberOfDuty = numberOfDuty
            i.save()
        # koniec liczenia

        person = Person.objects.order_by(Coalesce('numberOfDuty','surname').desc())     #lista wszystkich osob sortowana
        best = person[:5]   # pobranie tylko 5 pierwszych indeksow
        person = Person.objects.order_by(Coalesce('numberOfDuty', 'surname').asc()) # lista wszystkch sortowana asc
        worst = person[:5]  # pobranie tylko 5 pierwszych indeksow
        return render(request, 'index.html', {'best': best,'worst':worst,'all':personList})


def DetailsPage(request,id):
    person = Person.objects.get(idPerson=id)            #pobreanie osoby dla wyswietlenie szczegolow personalnych
    dutys = PersonOnDuty.objects.filter(idPerson=id)    #pobranie wszystych sluzb danej osoby dla tabeli
    return render(request, 'details.html', {'person':person,'dutys':dutys})


def DetailsPageCeremony(request,id):
    person = Person.objects.get(idPerson=id)    #pobreanie osoby dla wyswietlenie szczegolow personalnych
    dutys = PersonOnCeremony.objects.filter(idPerson=id)    #pobranie wszystych pedalow danej osoby dla tabeli
    return render(request, 'detailsCeremony.html', {'person':person,'dutys':dutys})

class CeremonyHomePage(TemplateView):
    def get(self, request, **kwargs):

        # pobranie listy wszystkich osob i policzenie ile laczenie maja godzin spedzonych na pedalowach
        personList = Person.objects.all()
        for i in personList:
            numberOfDuration = 0.0
            Ceremonys = PersonOnCeremony.objects.filter(idPerson=i.idPerson)
            for p in Ceremonys:
                numberOfDuration += p.idCeremony.duration
            i.numberOfCeremony = numberOfDuration
            i.save()
        # koniec listy osob

        personList.order_by(Coalesce('numberOfCeremony', 'surname').desc()) #sortowanie listy ze wzgledu na nazwiska
        return render(request, 'ceremony.html', {'all':personList})


class DashboardHomePage(TemplateView):
    def get(self, request, **kwargs):
        return allInOne(request) # funkcja zwracajaca witok gdzie sa przygotowane wszystkie formularze
                                # wywolywana po kazdym formularzu na dashboardzie do dodawania slub/pedalow/ludzi itp

def allInOne(request):
    personForm = PersonForm()
    dutyForm = DutyForm()
    personOnDutyForm = PersonOnDutyForm()
    ceremonyForm = CeremonyForm()
    personOnCeremonyForm = PersonOnCeremonyForm()
    personList = Person.objects.all()
    dutyList = Duty.objects.all().order_by(Coalesce('idDuty', 'date').desc())
    ceremonyList = Ceremony.objects.all().order_by(Coalesce('idCeremony', 'date').desc())
    return render(request, 'dashboard.html', {'personList': personList,
                                              'ceremonyList': ceremonyList,
                                              'dutyList': dutyList,
                                              'personForm': personForm,
                                              'dutyForm': dutyForm,
                                              'personOnDutyForm': personOnDutyForm,
                                              'ceremonyForm': ceremonyForm,
                                              'personOnCeremonyForm': personOnCeremonyForm
                                              })


def addPerson(request):
    if request.method == 'POST':                    #funkcja wywolana po formularzu dodawania osoby
        form = PersonForm(request.POST)
        if form.is_valid():
            person = form.save(commit=False)        # = nie zapisuj odrazu po sprawdzeniu walidacji
            person.numberOfDuty = 0                 #domyslnie ustaw sluzby na 0 i pedalowy na 0
            person.numberOfCeremony = 0
            person.save()
            return redirect('/dashboard')
        else:
            html = '<h1>Wystapil blad</h1>'
            return HttpResponse(html)

def addDuty(request):
    if request.method == 'POST':
        form = DutyForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/dashboard')
        else:
            html = '<h1>Wystapil blad</h1>'
            return HttpResponse(html)


def addCeremony(request):
    if request.method == 'POST':
        form = CeremonyForm(request.POST)
        if form.is_valid():
            form.save()

            return redirect('/dashboard')
        else:
            html = '<h1>Wystapil blad</h1>'
            return HttpResponse(html)

def addPersonToCeremony(request):
    if request.method == 'POST':
        form = PersonOnCeremonyForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/dashboard')
        else:
            html = '<h1>Wystapil blad</h1>'
            return HttpResponse(html)


def addToDuty(request):
    if request.method == 'POST':
        form = PersonOnDutyForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/dashboard')
        else:
            html = '<h1>Wystapil blad</h1>'
            return HttpResponse(html)


class GeneratorHomePage(TemplateView):                          #strona dla generatora PJ
    def get(self, request, **kwargs):
        all = Soldier.objects.all().order_by('idSoldier')       #wszyscy z kompanii sotrowani po id
                                                                #wszelkie zmiany nalezy robic poprzez usuniecie i
                                                                #zaimportowanie nowego excela z danymi
        return render(request, 'orderGenerate.html', {'all':all})


def generateOrderToWord(request):
        if request.method == 'POST':
            # tworzenie nowego dokumetu
            doc = Document()
            paragraph = doc.add_paragraph()
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(0)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(1.25)
                section.right_margin = Cm(1.5)

            makeHeadline(doc)

            uoList = request.POST.getlist('uo')
            unList = request.POST.getlist('un')

            ## sprawdzanie czy jest opcja ktoregos z urlopu zeby naglowek dodac do pliku
            if len(uoList) !=0 or len(unList) !=0:
                paragraph2 = doc.add_paragraph()
                paragraph2.style = doc.styles['Normal']
                paragraph2.add_run('5. URLOPY:').bold = True

            ### tworzenie listy na urlop okolicznosciowy ###
            if len(uoList) != 0:
                makeUo(request,uoList,doc)

            ### tworzenie listy na urlop nagrodowy ###
            if len(unList) != 0:
                makeUn(request,unList,doc)

            ### tworzenie listy HDK ###
            hdkList = request.POST.getlist('hdk')
            if len(hdkList) != 0:
                makeHDK(request,hdkList,doc)

            ### tworzenie PJ #####
            pjList = request.POST.getlist('pj')
            if len(pjList) != 0:
                makePJ(request,pjList,doc)

            # ustawienie aby Word był pobieralny ze strony
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=ListaPj.docx'
            doc.save(response)

        return response

def makeHeadline(doc):
    # przygotowanie nagłówka dokumentu aby obslugiwał poskie znaki
    headOfDoc = 'PODRÓŻE SŁUŻBOWE, URLOPY, ZWOLNIENIA OD ZAJĘĆ'
    convertHeadOfDoc = json.dumps(headOfDoc)
    finalHeadOfDoc = json.loads(convertHeadOfDoc)

    # stylizacja dla całego dokumentu
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # stylizacja tabeli #
    table_format = doc.styles['Table Grid']
    tableSpacing = table_format.paragraph_format
    # paragraph_format.line_spacing = 0
    tableSpacing.line_spacing = 1.1

    # dodawanie paragrafu
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    paragraph.add_run('V    ').bold = True
    runner = paragraph.add_run(finalHeadOfDoc)
    runner.bold = True
    runner.underline = True

    style = doc.styles['Body Text']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # footer
    footer = doc.sections[0].footer
    myFooter = footer.add_paragraph('DEWD 272/2/18/57/Batalion Szkolny/20'
                                    '                                 '
                                    '                                                                        ')
    myFooter.style = doc.styles['Body Text']
    runnerFooter = myFooter.add_run('JAWNE')
    runnerFooter.bold = True
    runnerFooter.underline = True

    numberOfPage = footer.add_paragraph()
    numberOfPage.add_run('1/1')
    numberOfPage.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # koniec footera

    #### dodawanie naglowka ####
    tommorow = date.today() + timedelta(days=1)
    dataOfOrder = tommorow.strftime("%d.%m.%Y")
    header = doc.sections[0].header
    p1 = header.add_paragraph()
    p1.add_run('Warszawa dn. ' + dataOfOrder + ' r.').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p2 = header.add_paragraph()
    p2.add_run('                             2 BATALION SZKOLNY \n'
               '             WOJSKOWA AKADEMIA TECHNICZNA                                                   ').bold = True
    runner = p2.add_run('JAWNE\n')
    runner.bold = True
    runner.underline = True

    headOfDoc2 = 'im. Jarosława Dąbrowskiego'
    convertHeadOfDoc2 = json.dumps(headOfDoc2)
    finalHeadOfDoc2 = json.loads(convertHeadOfDoc2)
    runner = p2.add_run('                                     ' + finalHeadOfDoc2 +
                        '                                                                      Egzemplarz pojedynczy')
    font2 = runner.font
    font2.name = 'Times New Roman'
    font2.size = Pt(10)
    runner.bold = True

    # koniec nagłowka


def makePJ(request,pjList,doc):
    personToPJ = []
    for i in pjList:
        personToPJ.append(Soldier.objects.get(idSoldier=i))

    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('7. PRZEPUSTKI JEDNORAZOWE').bold = True
    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('a) z 5 kp.').bold = True

    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    iterator = 1
    for i in personToPJ:
        row_cells = table.add_row().cells
        row_cells[0].text = str(iterator) + ')'
        row_cells[1].text = i.degree
        row_cells[2].text = i.name
        row_cells[3].text = i.surname
        row_cells[4].text = 'w dn.'

        if len(request.POST.get(str(i.idSoldier) + 'pj')) != 0:
            row_cells[5].text = request.POST.get(str(i.idSoldier) + 'pj')
            #row_cells[5].width = Inches(1.6)
        else:
            row_cells[5].text = request.POST.get('defaultDate')
            #row_cells[5].width = Inches(1.6)
        iterator = iterator + 1

        row_cells[0].width = Inches(0.2)
        row_cells[1].width = Inches(1.2)
        row_cells[3].width = Inches(1.6)
        row_cells[4].width = Inches(0.6)
        row_cells[5].width = Inches(1.1)

    paragraph = doc.add_paragraph()

def makeHDK(request,hdkList,doc):
    personToHDK = []
    for i in hdkList:
        personToHDK.append(Soldier.objects.get(idSoldier=i))

    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('6. ZWOLNIENIA').bold = True
    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('4) z tytułu honorowego krwiodastwa').bold = True
    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('a) z 5 kp.').bold = True

    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    iterator = 1
    for i in personToHDK:
        row_cells = table.add_row().cells
        row_cells[0].text = str(iterator) + ')'
        row_cells[1].text = i.degree
        row_cells[2].text = i.name
        row_cells[3].text = i.surname
        row_cells[4].text = 'w dn.'
        if len(request.POST.get(str(i.idSoldier) + 'hdk')) != 0:
            customData = request.POST.get(str(i.idSoldier) + 'hdk')
            splitData = customData.split('-')
            first = splitData[0]
            second = splitData[1]
            row_cells[5].text = first
            row_cells[6].text = '-'
            row_cells[7].text = second
        else:
            row_cells[5].text = request.POST.get('defaultDate')
            row_cells[6].text = '-'
            row_cells[7].text = request.POST.get('defaultDate')
        iterator = iterator + 1

        row_cells[0].width = Inches(0.2)
        row_cells[1].width = Inches(1.2)
        row_cells[2].width = Inches(1)
        row_cells[3].width = Inches(1.6)
        row_cells[4].width = Inches(0.6)
        row_cells[5].width = Inches(0.6)
        row_cells[6].width = Inches(0.2)

    paragraph = doc.add_paragraph()

def makeUn(request,unList,doc):
    personToUn = []
    for i in unList:
        personToUn.append(Soldier.objects.get(idSoldier=i))

    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('4) nagrodowy:').bold = True
    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('a) z 5 kp.').bold = True

    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    iterator = 1
    for i in personToUn:
        row_cells = table.add_row().cells
        row_cells[0].text = str(iterator) + ')'
        row_cells[1].text = i.degree
        row_cells[2].text = i.name
        row_cells[3].text = i.surname
        row_cells[4].text = 'w dn.'
        if len(request.POST.get(str(i.idSoldier) + 'un')) != 0:
            customData = request.POST.get(str(i.idSoldier) + 'un')
            splitData = customData.split('-')
            first = splitData[0]
            second = splitData[1]
            row_cells[5].text = first
            row_cells[6].text = '-'
            row_cells[7].text = second
        else:
            row_cells[5].text = request.POST.get('defaultDate')
            row_cells[6].text = '-'
            row_cells[7].text = request.POST.get('defaultDate')
        iterator = iterator + 1

        row_cells[0].width = Inches(0.2)
        row_cells[1].width = Inches(1.2)
        row_cells[2].width = Inches(1)
        row_cells[3].width = Inches(1.6)
        row_cells[4].width = Inches(0.6)
        row_cells[5].width = Inches(0.6)
        row_cells[6].width = Inches(0.2)

    paragraph = doc.add_paragraph()

def makeUo(request,uoList,doc):
    personToUo = []
    for i in uoList:
        personToUo.append(Soldier.objects.get(idSoldier=i))

    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('3) okolicznościowy').bold = True
    paragraph2 = doc.add_paragraph()
    paragraph2.style = doc.styles['Normal']
    paragraph2.add_run('a) z 5 kp.').bold = True

    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    iterator = 1
    for i in personToUo:
        row_cells = table.add_row().cells
        row_cells[0].text = str(iterator) + ')'
        row_cells[1].text = i.degree
        row_cells[2].text = i.name
        row_cells[3].text = i.surname
        row_cells[4].text = 'w dn.'
        if len(request.POST.get(str(i.idSoldier) + 'uo')) != 0:
            customData = request.POST.get(str(i.idSoldier) + 'uo')
            splitData = customData.split('-')
            first = splitData[0]
            second = splitData[1]
            row_cells[5].text = first
            row_cells[6].text = '-'
            row_cells[7].text = second
        else:
            row_cells[5].text = request.POST.get('defaultDate')
            row_cells[6].text = '-'
            row_cells[7].text = request.POST.get('defaultDate')
        iterator = iterator + 1

        row_cells[0].width = Inches(0.2)
        row_cells[1].width = Inches(1.2)
        row_cells[2].width = Inches(1)
        row_cells[3].width = Inches(1.6)
        row_cells[4].width = Inches(0.6)
        row_cells[5].width = Inches(0.6)
        row_cells[6].width = Inches(0.2)

    paragraph = doc.add_paragraph()
